from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.table import _Cell, Table

NBSP_CHARS = "\u00A0\u2007\u202F"


def replace_nbsp(text: str) -> str:
    for ch in NBSP_CHARS:
        text = text.replace(ch, " ")
    return text


@dataclass(frozen=True)
class LogicalCellKey:
    table_index: int
    row_index: int
    col_index: int


@dataclass
class LogicalCell:
    key: LogicalCellKey
    merged_rect: Tuple[int, int, int, int]  # (row_min, col_min, row_max_inclusive, col_max_inclusive)
    text: str
    # Index map from flattened text global index -> (paragraph_idx, run_idx, offset_in_run)
    index_map: List[Tuple[int, int, int]]


def load_document(path: str | Path) -> Document:
    return Document(str(path))


def iter_tables(doc: Document) -> Iterable[Tuple[int, Table]]:
    for t_idx, tbl in enumerate(doc.tables):
        yield t_idx, tbl


def _flatten_cell_with_index_map(cell: _Cell) -> Tuple[str, List[Tuple[int, int, int]]]:
    parts: List[str] = []
    index_map: List[Tuple[int, int, int]] = []
    for p_idx, para in enumerate(cell.paragraphs):
        for r_idx, run in enumerate(para.runs):
            txt = run.text or ""
            if not txt:
                continue
            norm = replace_nbsp(txt)
            parts.append(norm)
            for offset in range(len(norm)):
                index_map.append((p_idx, r_idx, offset))
    return ("".join(parts), index_map)


def extract_logical_cells(doc: Document) -> Tuple[List[LogicalCell], Dict[LogicalCellKey, _Cell]]:
    logical_cells: List[LogicalCell] = []
    key_to_cell: Dict[LogicalCellKey, _Cell] = {}

    for t_idx, tbl in iter_tables(doc):
        # Map each merged group (by underlying tc identity) to its first seen (row, col) owner
        gid_to_owner: Dict[int, Tuple[int, int]] = {}
        owner_to_slots: Dict[Tuple[int, int], List[Tuple[int, int]]] = {}

        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                gid = id(cell._tc)
                owner = gid_to_owner.setdefault(gid, (r_idx, c_idx))
                owner_to_slots.setdefault(owner, []).append((r_idx, c_idx))

        # For each owner cell, produce a logical cell
        for (owner_r, owner_c), slots in sorted(owner_to_slots.items()):
            owner_cell = tbl.cell(owner_r, owner_c)
            text, index_map = _flatten_cell_with_index_map(owner_cell)
            rows = [r for r, _ in slots]
            cols = [c for _, c in slots]
            rect = (min(rows), min(cols), max(rows), max(cols))
            key = LogicalCellKey(t_idx, owner_r, owner_c)
            logical_cells.append(
                LogicalCell(key=key, merged_rect=rect, text=text, index_map=index_map)
            )
            key_to_cell[key] = owner_cell

    logical_cells.sort(key=lambda lc: (lc.key.table_index, lc.key.row_index, lc.key.col_index))
    return logical_cells, key_to_cell