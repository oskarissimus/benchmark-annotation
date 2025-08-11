from __future__ import annotations

import shutil
import subprocess
from pathlib import Path
from typing import Dict, List, Sequence, Tuple

from docx import Document
from docx.shared import RGBColor

from .docx_utils import extract_logical_cells, load_document
from .evaluator import _map_eval_base_to_gt, _original_index_from_base_index


# White text on dark backgrounds
MISS_FONT = RGBColor(255, 255, 255)
MISS_FILL = "B00020"  # dark red background

MISPL_FONT = RGBColor(255, 255, 255)
MISPL_FILL = "0B6B00"  # dark green background


def _ensure_rPr(run):
    rPr = run._element.get_or_add_rPr()
    return rPr


def _set_run_shading(run, fill_hex: str) -> None:
    rPr = _ensure_rPr(run)
    # Add or update w:shd with fill color
    from docx.oxml.shared import OxmlElement, qn

    shd = rPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _clone_run_style(src_run, dst_run):
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    dst_run.font.color.rgb = src_run.font.color.rgb if src_run.font.color.rgb else None


def _insert_run_after(paragraph, ref_run, new_run):
    ref_run._element.addnext(new_run._element)


def _split_and_annotate_range(cell, index_map: List[Tuple[int, int, int]], start_idx: int, length: int, font_color: RGBColor, fill_hex: str, token_text: str | None = None):
    # Build ranges across runs
    end_idx = start_idx + length
    positions = list(range(start_idx, end_idx))
    if not positions:
        return

    # Group positions by (p_idx, r_idx)
    group: Dict[Tuple[int, int], List[int]] = {}
    for pos in positions:
        p_idx, r_idx, off = index_map[pos]
        group.setdefault((p_idx, r_idx), []).append(off)

    # Sort by paragraph and run
    grouped_items = sorted(group.items(), key=lambda kv: kv[0])

    # First segment determines insertion point
    first_p_idx, first_r_idx = grouped_items[0][0]
    paragraph = cell.paragraphs[first_p_idx]
    run = paragraph.runs[first_r_idx]

    # Collect token text from document if not provided
    collected_token: List[str] = []

    # Process from first to last run, removing token parts
    for (p_i, r_i), offs in grouped_items:
        para = cell.paragraphs[p_i]
        r = para.runs[r_i]
        offs.sort()
        start_off = offs[0]
        end_off = offs[-1] + 1
        txt = r.text
        before = txt[:start_off]
        inside = txt[start_off:end_off]
        after = txt[end_off:]
        collected_token.append(inside)
        # Replace current run text with before; append after as a new run inserted after
        r.text = before
        if after:
            after_run = para.add_run(after)
            _clone_run_style(r, after_run)
            _insert_run_after(para, r, after_run)

    final_token = token_text if token_text is not None else "".join(collected_token)

    # Insert the annotated token run after the first run
    token_run = paragraph.add_run(final_token)
    _clone_run_style(run, token_run)
    _insert_run_after(paragraph, run, token_run)
    token_run.font.color.rgb = font_color
    _set_run_shading(token_run, fill_hex)


def _insert_token_at_index(cell, index_map: List[Tuple[int, int, int]], insert_base_index: int, spans: Sequence[Tuple[int, int]], token_text: str, font_color: RGBColor, fill_hex: str):
    # Convert base index to original flattened index
    orig_index = _original_index_from_base_index(spans, insert_base_index)
    # Insert new run at this index; if it falls between runs, place at start of next run
    if orig_index >= len(index_map):
        # Append at end of last paragraph
        para = cell.paragraphs[-1]
        new_run = para.add_run(token_text)
        new_run.font.color.rgb = font_color
        _set_run_shading(new_run, fill_hex)
        return
    p_idx, r_idx, off = index_map[orig_index]
    para = cell.paragraphs[p_idx]
    r = para.runs[r_idx]
    txt = r.text
    before = txt[:off]
    after = txt[off:]
    r.text = before
    new_run = para.add_run(token_text)
    _insert_run_after(para, r, new_run)
    if after:
        after_run = para.add_run(after)
        _insert_run_after(para, new_run, after_run)
    new_run.font.color.rgb = font_color
    _set_run_shading(new_run, fill_hex)


def _render_preview_png(output_dir: Path, gt_cells, evaluation_cells) -> None:
    # Draw a simple preview: one line per cell with inline white text on dark red/green backgrounds
    from PIL import Image, ImageDraw, ImageFont

    lines = []  # type: List[List[tuple]]
    # Colors
    txt_color = (20, 20, 20)
    meta_color = (120, 120, 120)
    red_bg = (0xB0, 0x00, 0x20)      # dark red
    red_fg = (255, 255, 255)         # white text
    green_bg = (0x0B, 0x6B, 0x00)    # dark green
    green_fg = (255, 255, 255)       # white text

    for gt_cell, cell_result in zip(gt_cells, evaluation_cells):
        prefix = f"T{gt_cell.key.table_index} R{gt_cell.key.row_index} C{gt_cell.key.col_index}: "
        base_text = gt_cell.text
        gt_data = cell_result["gt_data"]
        ev_data = cell_result["ev_data"]

        # Map gt->ev positions
        from difflib import SequenceMatcher
        sm = SequenceMatcher(a=gt_data.base_text, b=ev_data.base_text, autojunk=False)
        opcodes = sm.get_opcodes()
        mapped_gt_positions = []
        for pos in gt_data.base_indices:
            ev_pos = 0
            for tag, i1, i2, j1, j2 in opcodes:
                if i1 <= pos <= i2:
                    if tag == "equal":
                        ev_pos = j1 + (pos - i1)
                    else:
                        ev_pos = j1
                    break
            mapped_gt_positions.append(ev_pos)
        ev_base_set = set(ev_data.base_indices)
        correct_positions = set(idx for idx in mapped_gt_positions if idx in ev_base_set)

        # Build events for rendering
        inserts: List[Tuple[int, str]] = []  # (orig_index, text)
        highlights: List[Tuple[int, int, Tuple[int,int,int], Tuple[int,int,int]]] = []  # (start, end, bg, fg)

        # Missed -> highlight original spans (dark red bg, white fg)
        for (start_idx, token_text), mapped in zip(gt_data.tokens, mapped_gt_positions):
            if mapped in correct_positions:
                continue
            highlights.append((start_idx, start_idx + len(token_text), red_bg, red_fg))

        # Misplaced -> insert token at mapped GT base position (dark green bg, white fg)
        for (ev_start, ev_token), ev_base_idx in zip(ev_data.tokens, ev_data.base_indices):
            if ev_base_idx in correct_positions:
                continue
            gt_base_pos = _map_eval_base_to_gt(gt_data.base_text, ev_data.base_text, [ev_base_idx])[0]
            orig_index = _original_index_from_base_index(gt_data.spans, gt_base_pos)
            inserts.append((orig_index, ev_token))

        # Apply inserts to compose display string and transform highlight indices accordingly
        inserts.sort(key=lambda x: x[0])
        display = base_text
        shift = 0
        for idx, text in inserts:
            pos = idx + shift
            display = display[:pos] + text + display[pos:]
            # shift highlights after pos
            for i, (s,e,bg,fg) in enumerate(highlights):
                if s >= idx:
                    highlights[i] = (s + len(text), e + len(text), bg, fg)
            shift += len(text)

        # Build segments for the line
        segments = []  # type: List[tuple]
        # Start with prefix as meta gray
        segments.append((prefix, None, meta_color))

        # Sort highlights by start
        highlights.sort(key=lambda x: x[0])
        cursor = 0
        for s, e, bg, fg in highlights:
            if s > cursor:
                segments.append((display[cursor:s], None, txt_color))
            segments.append((display[s:e], bg, fg))
            cursor = e
        if cursor < len(display):
            segments.append((display[cursor:], None, txt_color))

        lines.append(segments)

    # Layout image
    width = 1400
    line_height = 28
    padding = 16
    height = padding*2 + line_height * max(1, len(lines))
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("DejaVuSans.ttf", 16)
    except Exception:
        font = ImageFont.load_default()

    y = padding
    for segments in lines:
        x = padding
        for text, bg, fg in segments:
            if text:
                w = draw.textlength(text, font=font)
                if bg is not None:
                    draw.rounded_rectangle([x-2, y-4, x+w+2, y+line_height-10], radius=3, fill=bg)
                draw.text((x, y), text, fill=fg if fg else (20,20,20), font=font)
                x += int(w)
        y += line_height

    out_png = output_dir / "annotated-1.png"
    img.save(out_png)
    # Also write a PDF by saving the PNG as single-page PDF
    out_pdf = output_dir / "annotated.pdf"
    try:
        img.save(out_pdf, format="PDF")
    except Exception:
        pass


def generate_annotations(gt_path: str | Path, eval_path: str | Path, evaluation: Dict, output_dir: str | Path, debug: bool = False) -> None:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # Create a binary copy of GT to preserve formatting
    gt_copy_path = output_dir / "annotated.docx"
    shutil.copyfile(gt_path, gt_copy_path)

    # Open the copy for editing
    doc = Document(str(gt_copy_path))

    gt_cells, key_to_cell = extract_logical_cells(doc)

    # Iterate through evaluation per cell
    for cell_result, gt_cell in zip(evaluation["cells"], gt_cells):
        cell = key_to_cell[gt_cell.key]
        gt_data = cell_result["gt_data"]
        ev_data = cell_result["ev_data"]

        # Determine correctness set based on mapped positions
        from difflib import SequenceMatcher as _SM

        sm = _SM(a=gt_data.base_text, b=ev_data.base_text, autojunk=False)
        opcodes = sm.get_opcodes()
        mapped_gt_positions = []
        for pos in gt_data.base_indices:
            ev_pos = 0
            for tag, i1, i2, j1, j2 in opcodes:
                if i1 <= pos <= i2:
                    if tag == "equal":
                        ev_pos = j1 + (pos - i1)
                    elif tag in ("replace", "delete"):
                        ev_pos = j1
                    elif tag == "insert":
                        ev_pos = j1
                    break
            mapped_gt_positions.append(ev_pos)
        ev_base_set = set(ev_data.base_indices)
        correct_positions = set(idx for idx in mapped_gt_positions if idx in ev_base_set)

        # Apply dark red background with white text for missed
        for (start_idx, token_text), base_idx in zip(
            gt_data.tokens, gt_data.base_indices
        ):
            mapped = mapped_gt_positions[gt_data.base_indices.index(base_idx)]
            if mapped in correct_positions:
                continue
            _split_and_annotate_range(
                cell,
                gt_cell.index_map,
                start_idx,
                len(token_text),
                MISS_FONT,
                MISS_FILL,
            )

        # Misplaced: insert eval-only tokens mapped into GT base with dark green background and white text
        for (ev_start, ev_token), ev_base_idx in zip(ev_data.tokens, ev_data.base_indices):
            if ev_base_idx in correct_positions:
                continue
            gt_base_pos = _map_eval_base_to_gt(gt_data.base_text, ev_data.base_text, [ev_base_idx])[0]
            _insert_token_at_index(
                cell,
                gt_cell.index_map,
                gt_base_pos,
                gt_data.spans,
                ev_token,
                MISPL_FONT,
                MISPL_FILL,
            )

    # Save annotated doc
    doc.save(str(gt_copy_path))

    # Try converting to PDF via soffice; fallback to custom preview rendering if unavailable or pdf->png fails
    pdf_path = output_dir / "annotated.pdf"
    try:
        subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(gt_copy_path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    except Exception:
        _render_preview_png(output_dir, gt_cells, evaluation["cells"])
        return

    if not pdf_path.exists():
        # LibreOffice may emit a differently named PDF; attempt to locate
        for p in output_dir.glob("*.pdf"):
            pdf_path = p
            break

    # Convert PDF to PNGs
    try:
        from pdf2image import convert_from_path

        images = convert_from_path(str(pdf_path))
        for i, img in enumerate(images, start=1):
            img.save(output_dir / f"annotated-{i}.png")
    except Exception:
        _render_preview_png(output_dir, gt_cells, evaluation["cells"])