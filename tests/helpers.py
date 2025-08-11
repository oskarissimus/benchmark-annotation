from __future__ import annotations

from pathlib import Path
from typing import List, Optional, Sequence, Tuple

from docx import Document


def create_doc_with_table(
    tables: List[Tuple[int, int]],  # list of (rows, cols)
    borders: bool = True,
) -> Document:
    doc = Document()
    for rows, cols in tables:
        tbl = doc.add_table(rows=rows, cols=cols)
        if borders:
            tbl.style = "Table Grid"
    return doc


def set_cell_text(cell, text_parts: Sequence[str]) -> None:
    para = cell.paragraphs[0]
    # Clear existing content
    for r in list(para.runs):
        r.clear()
    para.text = ""
    for part in text_parts:
        para.add_run(part)


def merge_cells(table, r1, c1, r2, c2):
    top_left = table.cell(r1, c1)
    bottom_right = table.cell(r2, c2)
    top_left.merge(bottom_right)


def save_doc(doc: Document, path: str | Path) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))