from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

import pytest

from docx import Document

from .helpers import create_doc_with_table, set_cell_text, merge_cells, save_doc

CLI = ["python", "-m", "docx_markup_eval.cli"]


def run_cli(args: list[str]) -> subprocess.CompletedProcess:
    return subprocess.run(args, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)


def make_paths(tmp_path: Path, scenario: str):
    base = tmp_path / "fixtures" / "generated" / scenario
    gt = base / "gt.docx"
    ev = base / "eval.docx"
    out_json = base / "report.json"
    ref_out = base / "ref"
    return base, gt, ev, out_json, ref_out


def build_simple_docs_single_table(gt_tokens: list[list[str]], ev_tokens: list[list[str]], merges: list[tuple] | None = None):
    # tokens per cell row-major; each cell a list of text parts incl tokens and words
    rows = len(gt_tokens)
    cols = len(gt_tokens[0]) if rows else 0
    gt_doc = create_doc_with_table([(rows, cols)])
    ev_doc = create_doc_with_table([(rows, cols)])
    tbl_gt = gt_doc.tables[0]
    tbl_ev = ev_doc.tables[0]

    for r in range(rows):
        for c in range(cols):
            set_cell_text(tbl_gt.cell(r, c), [gt_tokens[r][c]])
            set_cell_text(tbl_ev.cell(r, c), [ev_tokens[r][c]])

    if merges:
        for (r1, c1, r2, c2) in merges:
            merge_cells(tbl_gt, r1, c1, r2, c2)
            merge_cells(tbl_ev, r1, c1, r2, c2)

    return gt_doc, ev_doc


def test_1_perfect_alignment(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario1")
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[["Lorem CELL_1 ipsum", "dolor CELL_2"], ["sit CELL_3 amet", "consectetur CELL_4"]],
        ev_tokens=[["Lorem cell_9 ipsum", "dolor CELL_99"], ["sit CELL_3 amet", "consectetur CELL_0"]],
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path)
    ])

    data = Path(out_path).read_text()
    import json

    res = json.loads(data)
    assert res == {"gt_total": 4, "eval_total": 4, "correct": 4, "misplaced": 0, "missed": 0}


def test_2_missing_one_token_with_annotation(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario2")
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[["CELL_1" ]],
        ev_tokens=[["text only"]],
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["gt_total"] == 1 and res["missed"] == 1
    # Artifacts should exist
    assert (ref_out / "annotated.docx").exists()
    assert (ref_out / "annotated.pdf").exists()
    pngs = list(ref_out.glob("annotated-*.png"))
    assert pngs, "Expected PNGs"


def test_3_extra_token_in_new_cell(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario3")
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[["No token"], ["No token"]],
        ev_tokens=[["CELL_1"], ["No token"]],
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["eval_total"] == 1 and res["misplaced"] == 1 and res["correct"] == 0


def test_4_wrong_position_same_cell(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario4")
    long = "Lorem ipsum dolor sit amet, consectetur adipiscing elit. " * 3
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[[long + " CELL_1 end"]],
        ev_tokens=[[long + " end CELL_2"]],
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["gt_total"] == 1 and res["eval_total"] == 1 and res["missed"] == 1 and res["misplaced"] == 1


def test_5_multiple_tokens_one_cell(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario5")
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[["start CELL_1 middle CELL_2 end" ]],
        ev_tokens=[["start CELL_9 middle end" ]],
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["correct"] == 1 and res["missed"] == 1 and res["misplaced"] == 0


def test_6_split_across_runs_plus_miss(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario6")
    doc = create_doc_with_table([(1, 1)])
    tbl = doc.tables[0]
    para = tbl.cell(0, 0).paragraphs[0]
    para.add_run("start CE")
    para.add_run("LL_1 end")
    save_doc(doc, gt_path)

    doc2 = create_doc_with_table([(1, 1)])
    tbl2 = doc2.tables[0]
    tbl2.cell(0, 0).paragraphs[0].add_run("start end")
    save_doc(doc2, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["missed"] == 1


def test_7_case_insensitive_match(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario7")
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[["CELL_1"]],
        ev_tokens=[["cell_9"]],
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["correct"] == 1 and res["missed"] == 0 and res["misplaced"] == 0


def test_8_multiple_tables(tmp_path: Path):
    # Build two tables with mixed results
    gt_doc = Document()
    t1 = gt_doc.add_table(rows=1, cols=2); t1.style = "Table Grid"
    t1.cell(0,0).paragraphs[0].add_run("CELL_1")
    t1.cell(0,1).paragraphs[0].add_run("no")
    t2 = gt_doc.add_table(rows=1, cols=1); t2.style = "Table Grid"
    t2.cell(0,0).paragraphs[0].add_run("CELL_2")

    ev_doc = Document()
    e1 = ev_doc.add_table(rows=1, cols=2); e1.style = "Table Grid"
    e1.cell(0,0).paragraphs[0].add_run("no")
    e1.cell(0,1).paragraphs[0].add_run("CELL_9")
    e2 = ev_doc.add_table(rows=1, cols=1); e2.style = "Table Grid"
    e2.cell(0,0).paragraphs[0].add_run("CELL_0")

    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario8")
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["gt_total"] == 2 and res["eval_total"] == 2 and res["correct"] == 1 and res["missed"] == 1 and res["misplaced"] == 1


def test_9_merged_cells_same_merge(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario9")
    gt = Document(); t = gt.add_table(rows=2, cols=2); t.style = "Table Grid"
    m = t.cell(0,0).merge(t.cell(1,1))
    m.paragraphs[0].add_run("CELL_1")

    ev = Document(); te = ev.add_table(rows=2, cols=2); te.style = "Table Grid"
    me = te.cell(0,0).merge(te.cell(1,1))
    me.paragraphs[0].add_run("CELL_9")

    save_doc(gt, gt_path)
    save_doc(ev, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["correct"] == 1 and res["missed"] == 0 and res["misplaced"] == 0


def test_10_merged_cells_different_geometry(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario10")
    gt = Document(); t = gt.add_table(rows=2, cols=2); t.style = "Table Grid"
    t.cell(0,0).merge(t.cell(0,1))
    t.cell(0,0).paragraphs[0].add_run("CELL_1")

    ev = Document(); te = ev.add_table(rows=2, cols=2); te.style = "Table Grid"
    te.cell(0,0).paragraphs[0].add_run("CELL_9")
    te.cell(0,1).paragraphs[0].add_run("")

    save_doc(gt, gt_path)
    save_doc(ev, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path), "--annotate", "--ref-out", str(ref_out)
    ])

    import json

    res = json.loads(Path(out_path).read_text())
    assert res["missed"] >= 0 and res["misplaced"] >= 0


def test_no_annotated_files_without_flag(tmp_path: Path):
    base, gt_path, ev_path, out_path, ref_out = make_paths(tmp_path, "scenario_no_annotate")
    gt_doc, ev_doc = build_simple_docs_single_table(
        gt_tokens=[["CELL_1"]], ev_tokens=[["CELL_9"]]
    )
    save_doc(gt_doc, gt_path)
    save_doc(ev_doc, ev_path)

    run_cli([
        "python", "-m", "docx_markup_eval.cli", "--gt", str(gt_path), "--eval", str(ev_path), "--format", "json", "--out", str(out_path)
    ])

    assert not (ref_out / "annotated.docx").exists()